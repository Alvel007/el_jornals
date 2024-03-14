import json
import locale
import os
from datetime import datetime, timedelta

import docx
import docx2pdf
import pythoncom
import pytz
from django.contrib.auth.decorators import login_required
from django.core.paginator import EmptyPage, Paginator
from django.db.models import Q
from django.http import HttpResponse, JsonResponse
from django.shortcuts import (HttpResponseRedirect, get_object_or_404,
                              redirect, render, reverse)
from django.template import RequestContext
from django.utils import timezone
from django.utils.decorators import method_decorator
from django.views.generic import FormView, ListView
from docx.enum.text import WD_UNDERLINE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor
from el_journals.settings import (NUMBER_ENTRIES_OP_LOG_PAGE,
                                  RETENTION_PERIOD_COMPLETED_RECORDS,
                                  SIGNAL_ON_REQUEST, TOTAL_VISIBLE_RECORDS_OPJ)
from op_journal.models import (AutocompleteOption, AutofillDispModel,
                               FileModelOPJ, MainPageOPJournal)
from powerline.forms import PowerLineForm
from powerline.models import AdmittingStaff, PowerLine, ThirdPartyDispatchers
from substation.models import Substation

from .forms import (AutofillDispForm, CommentOPJForm, MainPageOPJournalForm,
                    OPJournalForm)


@login_required(login_url='/login/')
def op_journal_detail(request, pk):
    op_journal_entry = get_object_or_404(MainPageOPJournal, pk=pk)
    return render(request,
                  'op_journal/op_journal_detail.html',
                  {'op_journal_entry': op_journal_entry}
                  )


@login_required(login_url='/login/')
def op_journal_edit(request, pk):
    op_journal_entry = get_object_or_404(MainPageOPJournal, pk=pk)
    user_has_operational_staff = request.user.operational_staff.filter(
        id=op_journal_entry.substation.id
    ).exists()
    if user_has_operational_staff:
        if op_journal_entry.entry_is_valid:
            if request.method == "POST":
                form = OPJournalForm(request.POST, instance=op_journal_entry)
                if form.is_valid():
                    op_journal_entry = form.save(commit=False)
                    if not op_journal_entry.entry_is_valid:
                        op_journal_entry.special_regime_introduced = False
                        op_journal_entry.emergency_event = False
                        op_journal_entry.short_circuit = False
                    op_journal_entry.save()
                    closing_entry = op_journal_entry.closing_entry.first()
                    if closing_entry:
                        op_journal_entry.closing_entry.clear()
                        closing_entry = None
                    return HttpResponseRedirect(
                        reverse(
                            'sub_op_journal',
                            kwargs={'substation_slug': (op_journal_entry
                                                        .substation.slug)}))
            else:
                initial_data = {'planned_completion_date': timezone.localtime(
                    op_journal_entry.planned_completion_date).strftime(
                        '%Y-%m-%dT%H:%M')}
                form = OPJournalForm(instance=op_journal_entry,
                                     initial=initial_data)
            return render(
                request,
                'op_journal/op_journal_edit.html',
                {'op_journal_entry': op_journal_entry, 'form': form})
        else:
            return redirect('op_journal_detail', pk=pk)
    else:
        return redirect('op_journal_detail', pk=pk)


@method_decorator(login_required(login_url='/login/'), name='dispatch')
class OpJournalView(ListView, FormView):
    template_name = 'op_journal/index.html'
    context_object_name = 'model_op_journal_data'
    paginate_by = NUMBER_ENTRIES_OP_LOG_PAGE
    extra_context = {
        'title': 'Инструкция по ведению электронного оперативного журнала'}
    form_class = MainPageOPJournalForm
    comment_form_class = CommentOPJForm
    form_VL = PowerLineForm

    def get(self, request, *args, **kwargs):
        self.object_list = None
        substation_slug = self.kwargs.get('substation_slug')
        if not substation_slug:
            return self.render_to_response(
                {'template': 'op_journal/manual.html'})
        return super().get(request, *args, **kwargs)

    def post(self, request, *args, **kwargs):
        self.object_list = self.get_queryset()
        if 'reset' in request.POST:
            return self.reset_search()
        form = self.get_form()
        if form.is_valid():
            return self.form_valid(form)
        else:
            return self.form_invalid(form)

    def reset_search(self):
        substation_slug = self.kwargs.get('substation_slug')
        return HttpResponseRedirect(
            reverse('sub_op_journal', args=[substation_slug]))

    def get_queryset(self):
        substation_slug = self.kwargs.get('substation_slug')
        query = self.request.GET.get('query')
        start_date = self.request.GET.get('start_date')
        end_date = self.request.GET.get('end_date')
        substation = Substation.objects.get(slug=substation_slug)
        queryset = MainPageOPJournal.objects.none()
        context = RequestContext(self.request)
        if substation in self.request.user.admin_opj.all() \
                or substation in self.request.user.administrative_staff.all() \
                or substation in self.request.user.operational_staff.all():
            queryset = MainPageOPJournal.objects.filter(
                substation__slug=substation_slug)
            if query:
                queryset = queryset.filter(text__icontains=query)
            if start_date and end_date:
                moscow_tz = pytz.timezone('Europe/Moscow')
                start_date = timezone.datetime.strptime(
                    start_date, "%Y-%m-%d").date()
                end_date = timezone.datetime.strptime(
                    end_date, "%Y-%m-%d").date()
                start_date = moscow_tz.localize(
                    timezone.datetime.combine(
                        start_date, timezone.datetime.min.time()))
                end_date = moscow_tz.localize(
                    timezone.datetime.combine(
                        end_date, timezone.datetime.max.time()))
                queryset = queryset.filter(Q(
                    pub_date__gte=start_date,
                    pub_date__lte=end_date))
            queryset = queryset.order_by(
                '-pub_date', '-id')[:TOTAL_VISIBLE_RECORDS_OPJ]
        else:
            context['substation_name'] = substation.name
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        substation_slug = self.kwargs.get('substation_slug', None)
        if substation_slug:
            substation = Substation.objects.get(slug=substation_slug)
            context['substation'] = substation
            if context['object_list'].exists():
                substation_name = context['object_list'][0].substation.name
                context['head'] = f"Оперативный журнал {substation_name}"
                context['title'] = f"Оперативный журнал {substation_name}"
                context['substation_slug'] = substation_slug
            else:
                context['head'] = "Оперативный журнал"
                context['title'] = "Оперативный журнал"
                context['substation_name'] = substation.name
            if substation in self.request.user.operational_staff.all():
                form = kwargs.get(
                    'form',
                    MainPageOPJournalForm(initial={
                        'substation': substation},
                        substation_slug=substation_slug))
                context['form'] = form
                context['has_permission'] = (substation in self.request.user
                                             .operational_staff.all())
            if substation in self.request.user.administrative_staff.all():
                comment_form = self.comment_form_class()
                context['comment_form'] = comment_form
            if substation.dispatch_point:
                autofill_disp_queryset = PowerLine.objects.filter(
                    for_CUS_dispatcher=substation)
                autofill_form = AutofillDispForm(
                    initial={'name': autofill_disp_queryset.values_list(
                        'name', flat=True)},
                    autofill_disp_queryset=autofill_disp_queryset)
                context['autofill_form'] = autofill_form
                disp_autofill = AutofillDispModel.objects.filter(
                    cus_dispatcher=substation)
                context['disp_autofill'] = disp_autofill
                print(disp_autofill)

                data_vl = []
                for power_line in autofill_disp_queryset:
                    dispatch_company = power_line.disp_manage
                    if dispatch_company:
                        dispatchers = ThirdPartyDispatchers.objects.filter(
                            disp_center=dispatch_company)
                        dispatchers_names = [
                            dispatcher.disp_name for dispatcher in dispatchers]
                    else:
                        dispatchers_names = []
                    admitting_pers = AdmittingStaff.objects.filter(
                        for_CUS_dispatcher=substation)
                    admitting_names = [
                        admitting.name for admitting in admitting_pers]
                    endings = [
                        ending.name for ending in power_line.ending.all()]
                    data_vl.append({'id': power_line.id,
                                    'name': power_line.name,
                                    'dispatchers': dispatchers_names,
                                    'admitting': admitting_names,
                                    'endings': endings})
                data_vl = json.dumps(data_vl, ensure_ascii=False)
                context['data_vl'] = data_vl
                dispatcher_records = MainPageOPJournal.objects.filter(
                    substation__in=substation.dispatcher_for.all(),
                    entry_is_valid=True,
                    withdrawal_for_repair=True,
                    closing_entry=None
                ).order_by('-pub_date', '-id')
                records_by_substation = {}
                for record in dispatcher_records:
                    substation_name = record.substation.name
                    if substation_name not in records_by_substation:
                        records_by_substation[substation_name] = []
                    records_by_substation[substation_name].append(record)
                context['records_by_substation'] = records_by_substation
                now = timezone.now()
                for record in dispatcher_records:
                    time_difference = record.planned_completion_date - now
                    record.is_close_to_completion = (
                        (time_difference
                         .total_seconds()) < SIGNAL_ON_REQUEST * 3600)
                    record.is_past_due = time_difference.total_seconds() < 0
                context['records_by_substation'] = records_by_substation
                disp_work_records = MainPageOPJournal.objects.filter(
                    substation__in=substation.dispatcher_for.all(),
                    entry_is_valid=True,
                    permission_to_work=True,
                    closing_entry=None
                ).order_by('-pub_date', '-id')
                work_by_substation = {}
                for record in disp_work_records:
                    substation_name = record.substation.name
                    if substation_name not in work_by_substation:
                        work_by_substation[substation_name] = []
                    work_by_substation[substation_name].append(record)
                context['work_by_substation'] = work_by_substation

                completed_records_by_substation = {}
                for dispatcher in substation.dispatcher_for.all():
                    completed_records = MainPageOPJournal.objects.filter(
                        substation=dispatcher,
                        entry_is_valid=True,
                        withdrawal_for_repair=True
                    ).filter(
                        Q(closing_entry__isnull=False) |
                        Q(closing_entry__pub_date__lte=datetime.now() -
                          timedelta(days=RETENTION_PERIOD_COMPLETED_RECORDS))
                    ).order_by('-closing_entry__pub_date', '-id')
                    if completed_records.exists():
                        completed_records_by_substation[
                            dispatcher.name
                        ] = completed_records
                context[
                    'completed_records_by_substation'
                ] = completed_records_by_substation

                compl_worked_records_by_substation = {}
                for dispatcher in substation.dispatcher_for.all():
                    compl_worked_records = MainPageOPJournal.objects.filter(
                        substation=dispatcher,
                        entry_is_valid=True,
                        permission_to_work=True
                    ).filter(
                        Q(closing_entry__isnull=False) |
                        Q(closing_entry__pub_date__lte=datetime.now() -
                          timedelta(days=RETENTION_PERIOD_COMPLETED_RECORDS))
                    ).order_by('-closing_entry__pub_date', '-id')
                    if compl_worked_records.exists():
                        compl_worked_records_by_substation[
                            dispatcher.name
                        ] = compl_worked_records
                context[
                    'compl_worked_records_by_substation'
                ] = compl_worked_records_by_substation

            filter_autocomplite = AutocompleteOption.objects.filter(
                substation=substation)
            context['filter_autocomplite'] = filter_autocomplite
            filtered_records = MainPageOPJournal.objects.filter(
                substation=substation,
                entry_is_valid=True,
                withdrawal_for_repair=True,
                closing_entry=None
            ).order_by('-pub_date', '-id')
            now = timezone.now()
            for record in filtered_records:
                time_difference = record.planned_completion_date - now
                record.is_close_to_completion = (
                    time_difference.total_seconds() < SIGNAL_ON_REQUEST * 3600)
                record.is_past_due = time_difference.total_seconds() < 0
            context['filtered_model_op_journal_data'] = filtered_records

            works_records = MainPageOPJournal.objects.filter(
                substation=substation,
                entry_is_valid=True,
                permission_to_work=True,
                closing_entry=None
            ).order_by('-pub_date', '-id')
            context['works_model_op_journal_data'] = works_records

            disabling_records = MainPageOPJournal.objects.filter(
                substation=substation,
                entry_is_valid=True,
                withdrawal_for_repair=True
            ).filter(
                Q(closing_entry__isnull=False) |
                Q(closing_entry__pub_date__lte=datetime.now() -
                  timedelta(days=RETENTION_PERIOD_COMPLETED_RECORDS))
            ).order_by('-closing_entry__pub_date', '-id')
            context['disabling_model_op_journal_data'] = disabling_records

            worked_records = MainPageOPJournal.objects.filter(
                substation=substation,
                entry_is_valid=True,
                permission_to_work=True
            ).filter(
                Q(closing_entry__isnull=False) |
                Q(closing_entry__pub_date__lte=datetime.now() -
                  timedelta(days=RETENTION_PERIOD_COMPLETED_RECORDS))
            ).order_by('-closing_entry__pub_date', '-id')
            context['worked_model_op_journal_data'] = worked_records
        else:
            context['substation_name'] = None

        if 'object_list' in context and not context['object_list'].exists():
            paginator = Paginator(context['object_list'], self.paginate_by)
            last_page = paginator.num_pages
            page = self.request.GET.get('page', last_page)
            try:
                page = paginator.page(page)
            except EmptyPage:
                page = paginator.page(last_page)
            context['page_obj'] = page

        context[
            'RETENTION_PERIOD_COMPLETED_RECORDS'
        ] = RETENTION_PERIOD_COMPLETED_RECORDS
        return context

    def form_invalid(self, form):
        substation_slug = self.kwargs.get('substation_slug', None)
        if substation_slug:
            substation = Substation.objects.get(slug=substation_slug)
            form.fields[
                'existing_entry'
            ].queryset = MainPageOPJournal.objects.filter(
                entry_is_valid=True,
                withdrawal_for_repair=True,
                closing_entry=None,
                substation=substation)
            form.fields[
                'work_entry'
            ].queryset = MainPageOPJournal.objects.filter(
                entry_is_valid=True,
                permission_to_work=True,
                closing_entry=None,
                substation=substation)
        form.fields[
            'important_event_checkbox'
        ].initial = self.request.POST.get(
            'important_event_checkbox',
            False)
        return self.render_to_response(self.get_context_data(form=form))

    def form_valid(self, form):
        user = self.request.user
        form.instance.user = user
        op_journal = form.save(commit=False)
        important_event_checkbox = self.request.POST.get(
            'important_event_checkbox',
            False)
        if important_event_checkbox:
            op_journal.withdrawal_for_repair = True
        permission_to_work_checkbox = self.request.POST.get(
            'permission_to_work_checkbox',
            False)
        if permission_to_work_checkbox:
            op_journal.permission_to_work = True
        op_journal.save()
        existing_entry_id = self.request.POST.get('existing_entry')
        if existing_entry_id:
            existing_entry = get_object_or_404(
                MainPageOPJournal,
                id=existing_entry_id)
            existing_entry.save()
            op_journal.closing_entry.add(existing_entry)
        work_entry_id = self.request.POST.get('work_entry')
        if work_entry_id:
            work_entry = get_object_or_404(
                MainPageOPJournal,
                id=work_entry_id)
            work_entry.save()
            op_journal.closing_entry.add(work_entry)
        files = self.request.FILES.getlist('file')
        for file in files:
            file_instance = FileModelOPJ(
                main_page_op_journal=op_journal,
                file=file)
            file_instance.save()
        form.save_m2m()
        form.save()
        return HttpResponseRedirect(
            reverse(
                'sub_op_journal',
                args=[form.instance.substation.slug]))


@login_required(login_url='/login/')
def autocomplete_view(request, substation_slug):
    substation = get_object_or_404(Substation, slug=substation_slug)
    query = request.GET.get('term', '')
    options = AutocompleteOption.objects.filter(
        substation=substation,
        text__icontains=query)
    results = [{'label': option.label,
                'value': option.text,
                'out_of_work': option.out_of_work,
                'getting_started': option.getting_started,
                'disabling': option.disabling,
                'enabling': option.enabling
                } for option in options]
    if substation.dispatch_point:
        CHIO = AutofillDispModel.objects.filter(cus_dispatcher=substation)

        ADDITIONAL_ENTRY_ONE = [{
            'label': (option.hand_over_text[:140] + '...'),
            'value': option.hand_over_text,
            'out_of_work': False,
            'getting_started': False,
            'disabling': False,
            'enabling': False} for option in CHIO]
        results.extend(ADDITIONAL_ENTRY_ONE)
        ADDITIONAL_ENTRY_TWO = [{
            'label': (option.prm_tolerances[:140] + '...'),
            'value': option.prm_tolerances,
            'out_of_work': False,
            'getting_started': False,
            'disabling': True,
            'enabling': False} for option in CHIO]
        results.extend(ADDITIONAL_ENTRY_TWO)
        ADDITIONAL_ENTRY_THREE = [{
            'label': (option.prm_only[:140] + '...'),
            'value': option.prm_only,
            'out_of_work': False,
            'getting_started': False,
            'disabling': False,
            'enabling': False} for option in CHIO]
        results.extend(ADDITIONAL_ENTRY_THREE)
        ADDITIONAL_ENTRY_FOUR = [{
            'label': (option.admission_omly[:140] + '...'),
            'value': option.admission_omly,
            'out_of_work': False,
            'getting_started': False,
            'disabling': True,
            'enabling': False} for option in CHIO]
        results.extend(ADDITIONAL_ENTRY_FOUR)
        ADDITIONAL_ENTRY_FIVE = [{
            'label': (option.without_tripping[:140] + '...'),
            'value': option.without_tripping,
            'out_of_work': False,
            'getting_started': False,
            'disabling': True,
            'enabling': False} for option in CHIO]
        results.extend(ADDITIONAL_ENTRY_FIVE)
        ADDITIONAL_ENTRY_SIX = [{
            'label': (option.at_substation[:140] + '...'),
            'value': option.at_substation,
            'out_of_work': False,
            'getting_started': False,
            'disabling': True,
            'enabling': False} for option in CHIO]
        results.extend(ADDITIONAL_ENTRY_SIX)
        ADDITIONAL_ENTRY_SEVEN = [{
            'label': (option.and_work[:140] + '...'),
            'value': option.and_work,
            'out_of_work': False,
            'getting_started': True,
            'disabling': False,
            'enabling': False} for option in CHIO]
        results.extend(ADDITIONAL_ENTRY_SEVEN)
        ADDITIONAL_ENTRY_EIGHT = [{
            'label': (option.submit_vl[:140] + '...'),
            'value': option.submit_vl,
            'out_of_work': False,
            'getting_started': False,
            'disabling': False,
            'enabling': False} for option in CHIO]
        results.extend([ADDITIONAL_ENTRY_EIGHT])

    return JsonResponse(results, safe=False)


@login_required(login_url='/login/')
def add_comment(request, post_id):
    post = get_object_or_404(MainPageOPJournal, pk=post_id)

    if request.method == 'POST':
        comment_form = CommentOPJForm(request.POST)
        if comment_form.is_valid():
            comment = comment_form.save(commit=False)
            comment.user = request.user
            comment.save()
            post.comment = comment
            post.save()
    return HttpResponseRedirect(
        reverse(
            'sub_op_journal',
            args=[post.substation.slug]))


@login_required(login_url='/login/')
def export_records(request, substation_slug):
    start_date_str = request.POST.get('start_date')
    end_date_str = request.POST.get('end_date')
    start_date = timezone.make_aware(
        datetime.strptime(start_date_str, '%Y-%m-%d'))
    end_date = timezone.make_aware(
        datetime.strptime(end_date_str, '%Y-%m-%d'))
    end_date += timedelta(days=1)

    local_timezone = pytz.timezone('Europe/Moscow')

    queryset = MainPageOPJournal.objects.filter(
        pub_date__range=[start_date.astimezone(local_timezone),
                         end_date.astimezone(local_timezone)],
        substation__slug=substation_slug)
    template = docx.Document('templates/op_journal/template.docx')
    substation_name = Substation.objects.get(slug=substation_slug).name
    for paragraph in template.paragraphs:
        if 'PLACEHOLDER_FOR_SUBSTATION_NAME' in paragraph.text:
            for run in paragraph.runs:
                if 'PLACEHOLDER_FOR_SUBSTATION_NAME' in run.text:
                    run.text = run.text.replace(
                        'PLACEHOLDER_FOR_SUBSTATION_NAME',
                        substation_name)
            break
    table = template.tables[0]
    for record in queryset:
        row = table.add_row().cells
        local_time = record.pub_date.astimezone(local_timezone)

        if record.emergency_event and not record.short_circuit:
            cell = row[0]
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FF0000"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            paragraph = row[0].paragraphs[0]
            run = paragraph.add_run(local_time.strftime('%d.%m.%Y\n%H:%M'))
            font = run.font
            font.color.rgb = RGBColor(255, 255, 255)  # Белый цвет текста
            font.underline = True
            run.underline = WD_UNDERLINE.WAVY
        elif not record.emergency_event and record.short_circuit:
            cell = row[0]
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0000FF"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            paragraph = row[0].paragraphs[0]
            run = paragraph.add_run(local_time.strftime('%d.%m.%Y\n%H:%M'))
            font = run.font
            font.color.rgb = RGBColor(255, 255, 255)  # Белый цвет текста
            font.underline = True
            run.underline = WD_UNDERLINE.WAVY
        elif record.emergency_event and record.short_circuit:
            cell = row[0]
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="ffff00"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            paragraph = row[0].paragraphs[0]
            run = paragraph.add_run(local_time.strftime('%d.%m.%Y\n%H:%M'))
            font = run.font
            font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет текста
            font.underline = True
            run.underline = WD_UNDERLINE.WAVY
        else:
            row[0].text = local_time.strftime('%d.%m.%Y\n%H:%M')

        if record.entry_is_valid:
            if record.special_regime_introduced:
                paragraph = row[1].paragraphs[0]
                run = paragraph.add_run(record.text)
                font = run.font
                font.color.rgb = RGBColor(255, 0, 0)
                paragraph.add_run(f'\n{record.user.position} {record.user}')
            else:
                row[1].text = (f'{record.text}\n'
                               f'{record.user.position} {record.user}')
        else:
            paragraph = row[1].paragraphs[0]
            run = paragraph.add_run('ОШИБОЧНАЯ ЗАПИСЬ!\n')
            font = run.font
            font.strike = False
            run = paragraph.add_run(record.text)
            font = run.font
            font.strike = True
            paragraph.add_run(f'\n{record.user.position} {record.user}')

        row[2].text = ((f'{record.comment.text}\n'
                        f'{record.comment.user.position} '
                        f'{record.comment.user}') if record.comment else "")
    pythoncom.CoInitialize()
    temp_docx_path = 'temp.docx'
    template.save(temp_docx_path)
    temp_pdf_path = 'temp.pdf'
    docx2pdf.convert(temp_docx_path, temp_pdf_path)
    with open(temp_pdf_path, 'rb') as f:
        pdf_content = f.read()
    os.remove(temp_docx_path)
    os.remove(temp_pdf_path)
    response = HttpResponse(content_type='application/pdf')
    response[
        'Content-Disposition'
    ] = 'attachment; filename=exported_records.pdf'
    response.write(pdf_content)

    return response


def autofill_form_view(request, substation_slug):
    locale.setlocale(locale.LC_ALL, 'en_US.UTF-8')
    substation = Substation.objects.get(slug=substation_slug)
    if request.method == 'POST':
        print(request.POST)
        autofill_form = AutofillDispForm(request.POST)
        if autofill_form.is_valid():
            name = autofill_form.cleaned_data['name']
            dispatcher = autofill_form.cleaned_data['dispatcher']
            admitting = autofill_form.cleaned_data['admitting']
            cus_dispatcher = autofill_form.cleaned_data['cus_dispatcher']
            ending = autofill_form.cleaned_data['ending']
            if AutofillDispModel.objects.filter(
                    cus_dispatcher=cus_dispatcher).exists():
                autofill_instance = AutofillDispModel.objects.get(
                    cus_dispatcher=cus_dispatcher)
                autofill_instance.name = name
                autofill_instance.dispatcher = dispatcher
                autofill_instance.admitting = admitting
                autofill_instance.ending = ending
                autofill_instance.end_time = (
                    autofill_form.cleaned_data['end_time'])
                autofill_instance.emergency_entry = (
                    autofill_form.cleaned_data['emergency_entry'])
                autofill_instance.cus_dispatcher = substation.name
                autofill_instance.save()
                return HttpResponseRedirect(
                    reverse(
                        'sub_op_journal',
                        args=[substation.slug]))
            else:

                new_autofill_instance = autofill_form.save(commit=False)
                new_autofill_instance.name = name
                new_autofill_instance.dispatcher = dispatcher
                new_autofill_instance.ending = ending
                new_autofill_instance.admitting = admitting
                new_autofill_instance.cus_dispatcher = substation.name
                new_autofill_instance.save()
                return HttpResponseRedirect(
                    reverse(
                        'sub_op_journal',
                        args=[substation.slug]))
        else:
            print("Ошибка валидации")
    return HttpResponseRedirect(
        reverse(
            'sub_op_journal',
            args=[substation.slug]))

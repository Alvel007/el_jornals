from io import BytesIO
import pytz
import os
from django.conf import settings
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.utils import ImageReader
from reportlab.lib import pdfencrypt

from docx import Document
from django.template import Context
from docxtpl import DocxTemplate
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Inches, RGBColor
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
import io
import docx
import codecs
import pytz
import docx2pdf
import pythoncom


from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.shortcuts import render, redirect, reverse
from django.views.generic import ListView, FormView, TemplateView
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.utils.dateparse import parse_date
from django.http import JsonResponse, HttpResponse
from django.shortcuts import render, get_object_or_404, HttpResponseRedirect, redirect
from django.db.models import Q
from django.template.loader import get_template, render_to_string
from django.core.files.uploadedfile import UploadedFile
from django.template import RequestContext
from xhtml2pdf import pisa
from datetime import datetime, timedelta
from django.utils import timezone
from django.contrib import messages


from op_journal.models import MainPageOPJournal, AutocompleteOption, FileModelOPJ
from el_journals.settings import NUMBER_ENTRIES_OP_LOG_PAGE, STATICFILES_DIRS, STATIC_URL, MEDIA_URL, MEDIA_ROOT, TOTAL_VISIBLE_RECORDS_OPJ, RETENTION_PERIOD_COMPLETED_RECORDS
from .forms import MainPageOPJournalForm, OPJournalForm, CommentOPJForm
from substation.models import Substation

import PyPDF2
from PyPDF2 import PdfFileWriter


@login_required(login_url='/login/')
def op_journal_detail(request, pk):
    op_journal_entry = get_object_or_404(MainPageOPJournal, pk=pk)
    return render(request, 'op_journal/op_journal_detail.html', {'op_journal_entry': op_journal_entry})


@login_required(login_url='/login/')
def op_journal_edit(request, pk):
    op_journal_entry = get_object_or_404(MainPageOPJournal, pk=pk)
    if request.user == op_journal_entry.user:
        if op_journal_entry.entry_is_valid:
            if request.method == "POST":
                form = OPJournalForm(request.POST, instance=op_journal_entry)
                if form.is_valid():
                    op_journal_entry = form.save(commit=False)
                    if not op_journal_entry.entry_is_valid:
                        op_journal_entry.special_regime_introduced = False
                        op_journal_entry.emergency_event = False
                        op_journal_entry.short_circuit = False
                    op_journal_entry.save()
                    closing_entry = op_journal_entry.closing_entry.first()
                    if closing_entry:
                        closing_entry.important_event_date_over = None
                        closing_entry.save()
                    op_journal_entry.closing_entry.clear()
                    
                    return HttpResponseRedirect(reverse('sub_op_journal', kwargs={'substation_slug': op_journal_entry.substation.slug}))
            else:
                form = OPJournalForm(instance=op_journal_entry)
            return render(request, 'op_journal/op_journal_edit.html', {'op_journal_entry': op_journal_entry, 'form': form})
        else:
            return redirect('op_journal_detail', pk=pk)
    else:
        return redirect('op_journal_detail', pk=pk)

@method_decorator(login_required(login_url='/login/'), name='dispatch') 
class OpJournalView(ListView, FormView):
    template_name = 'op_journal/index.html'
    context_object_name = 'model_op_journal_data'
    paginate_by = NUMBER_ENTRIES_OP_LOG_PAGE
    extra_context = {'title': 'Инструкция по ведению электронного оперативного журнала'}
    form_class = MainPageOPJournalForm
    comment_form_class = CommentOPJForm

    def get(self, request, *args, **kwargs):
        self.object_list = None
        substation_slug = self.kwargs.get('substation_slug')
        if not substation_slug:
            return self.render_to_response({'template': 'op_journal/manual.html'})
        return super().get(request, *args, **kwargs)

    def post(self, request, *args, **kwargs):
        self.object_list = self.get_queryset()
        if 'reset' in request.POST:
            return self.reset_search()
        return super().post(request, *args, **kwargs)

    def reset_search(self):
        substation_slug = self.kwargs.get('substation_slug')
        return HttpResponseRedirect(reverse('sub_op_journal', args=[substation_slug]))

    def get_queryset(self): 
        substation_slug = self.kwargs.get('substation_slug') 
        query = self.request.GET.get('query') 
        start_date = self.request.GET.get('start_date') 
        end_date = self.request.GET.get('end_date') 
        substation = Substation.objects.get(slug=substation_slug) 
        queryset = MainPageOPJournal.objects.none() 
        context = RequestContext(self.request) 
        if substation in self.request.user.admin_opj.all() or substation in self.request.user.administrative_staff.all() or substation in self.request.user.operational_staff.all(): 
            queryset = MainPageOPJournal.objects.filter(substation__slug=substation_slug) 
            if query: 
                queryset = queryset.filter(text__icontains=query) 
            if start_date and end_date: 
                moscow_tz = pytz.timezone('Europe/Moscow') 
                start_date = timezone.datetime.strptime(start_date, "%Y-%m-%d").date() 
                end_date = timezone.datetime.strptime(end_date, "%Y-%m-%d").date() 
                start_date = moscow_tz.localize(timezone.datetime.combine(start_date, timezone.datetime.min.time())) 
                end_date = moscow_tz.localize(timezone.datetime.combine(end_date, timezone.datetime.max.time())) 
                queryset = queryset.filter(Q(pub_date__gte=start_date, pub_date__lte=end_date)) 
            queryset = queryset.order_by('-pub_date', '-id')[:TOTAL_VISIBLE_RECORDS_OPJ] 
        else: 
            context['substation_name'] = substation.name 
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        substation_slug = self.kwargs.get('substation_slug', None)
        if substation_slug:
            substation = Substation.objects.get(slug=substation_slug)
            context['substation'] = substation
            if context['object_list'].exists():
                substation_name = context['object_list'][0].substation.name
                context['head'] = f"Оперативный журнал {substation_name}"
                context['title'] = f"Оперативный журнал {substation_name}"
                context['substation_slug'] = substation_slug
            else:
                context['head'] = "Оперативный журнал"
                context['title'] = "Оперативный журнал"
                context['substation_name'] = substation.name
            if substation in self.request.user.operational_staff.all():
                form = kwargs.get('form', MainPageOPJournalForm(initial={'substation': substation}, substation_slug=substation_slug))
                context['form'] = form
                context['has_permission'] = substation in self.request.user.operational_staff.all()
            if substation in self.request.user.administrative_staff.all():
                comment_form = self.comment_form_class()
                context['comment_form'] = comment_form
            if substation.dispatch_point:
                dispatcher_records = MainPageOPJournal.objects.filter(
                    substation__in=substation.dispatcher_for.all(),
                    entry_is_valid=True,
                    important_event_date_start__isnull=False,
                    important_event_date_over__isnull=True
                ).order_by('-pub_date', '-id')
                records_by_substation = {}
                for record in dispatcher_records:
                    substation_name = record.substation.name
                    if substation_name not in records_by_substation:
                        records_by_substation[substation_name] = []
                    records_by_substation[substation_name].append(record)
                context['records_by_substation'] = records_by_substation

                completed_records_by_substation = {}
                for dispatcher in substation.dispatcher_for.all():
                    completed_records = MainPageOPJournal.objects.filter(
                        substation=dispatcher,
                        entry_is_valid=True,
                        important_event_date_start__isnull=False,
                        important_event_date_over__isnull=False,
                        important_event_date_over__gte=datetime.now() - timedelta(days=RETENTION_PERIOD_COMPLETED_RECORDS)
                    ).order_by('-important_event_date_start', '-id')
                    if completed_records.exists():
                        completed_records_by_substation[dispatcher.name] = completed_records
                context['completed_records_by_substation'] = completed_records_by_substation

            filtered_records = MainPageOPJournal.objects.filter(
                substation=substation,
                entry_is_valid=True,
                important_event_date_start__isnull=False,
                important_event_date_over__isnull=True
            ).order_by('-pub_date', '-id')
            context['filtered_model_op_journal_data'] = filtered_records

            disabling_records = MainPageOPJournal.objects.filter(
                substation=substation,
                entry_is_valid=True,
                important_event_date_start__isnull=False,
                important_event_date_over__isnull=False,
                important_event_date_over__gte=datetime.now() - timedelta(days=RETENTION_PERIOD_COMPLETED_RECORDS)
            ).order_by('-important_event_date_start', '-id')
            context['disabling_model_op_journal_data'] = disabling_records
        else:
            context['substation_name'] = None

        if 'object_list' in context and not context['object_list'].exists():
            paginator = Paginator(context['object_list'], self.paginate_by)
            last_page = paginator.num_pages
            page = self.request.GET.get('page', last_page)
            try:
                page = paginator.page(page)
            except EmptyPage:
                page = paginator.page(last_page)
            context['page_obj'] = page

        context['RETENTION_PERIOD_COMPLETED_RECORDS'] = RETENTION_PERIOD_COMPLETED_RECORDS
        return context


    def form_invalid(self, form):
        substation_slug = self.kwargs.get('substation_slug', None)
        if substation_slug:
            substation = Substation.objects.get(slug=substation_slug)
            form.fields['existing_entry'].queryset = MainPageOPJournal.objects.filter(
                entry_is_valid=True,
                important_event_date_start__isnull=False,
                important_event_date_over__isnull=True,
                substation=substation
            )
        return self.render_to_response(self.get_context_data(form=form))

    def form_valid(self, form):
        user = self.request.user
        form.instance.user = user
        op_journal = form.save(commit=False)
        if self.request.POST.get('important_event_checkbox'):
            op_journal.important_event_date_start = op_journal.pub_date
        op_journal.save()
        existing_entry_id = self.request.POST.get('existing_entry')
        if existing_entry_id:
            existing_entry = get_object_or_404(MainPageOPJournal, id=existing_entry_id)
            existing_entry.important_event_date_over = op_journal.pub_date
            existing_entry.save()
            op_journal.closing_entry.add(existing_entry)  # Добавляем связь между экземплярами
        files = self.request.FILES.getlist('file')
        for file in files:
            file_instance = FileModelOPJ(main_page_op_journal=op_journal, file=file)
            file_instance.save()
        form.save_m2m()
        return HttpResponseRedirect(reverse('sub_op_journal', args=[form.instance.substation.slug]))


@login_required(login_url='/login/')
def autocomplete_view(request, substation_slug):
    substation = get_object_or_404(Substation, slug=substation_slug)
    query = request.GET.get('term', '')
    options = AutocompleteOption.objects.filter(substation=substation, text__icontains=query)
    results = [option.text for option in options]
    return JsonResponse(results, safe=False)


@login_required(login_url='/login/')
def add_comment(request, post_id):
    post = get_object_or_404(MainPageOPJournal, pk=post_id)

    if request.method == 'POST':
        comment_form = CommentOPJForm(request.POST)
        if comment_form.is_valid():
            comment = comment_form.save(commit=False)
            comment.user = request.user
            comment.save()
            post.comment = comment
            post.save()
    return HttpResponseRedirect(reverse('sub_op_journal', args=[post.substation.slug]))


@login_required(login_url='/login/')
def export_records(request, substation_slug):
    start_date_str = request.POST.get('start_date')
    end_date_str = request.POST.get('end_date')
    start_date = timezone.make_aware(datetime.strptime(start_date_str, '%Y-%m-%d'))
    end_date = timezone.make_aware(datetime.strptime(end_date_str, '%Y-%m-%d'))
    end_date += timedelta(days=1)

        # Получение локальной временной зоны сервера
    local_timezone = pytz.timezone('Europe/Moscow')  # Замените 'Europe/Moscow' на соответствующую временную зону

    queryset = MainPageOPJournal.objects.filter(pub_date__range=[start_date.astimezone(local_timezone), end_date.astimezone(local_timezone)], substation__slug=substation_slug)

        # Загрузка шаблона Word
    template = docx.Document('templates/op_journal/template.docx')
    substation_name = Substation.objects.get(slug=substation_slug).name
    for paragraph in template.paragraphs:
        if 'PLACEHOLDER_FOR_SUBSTATION_NAME' in paragraph.text:
            for run in paragraph.runs:
                if 'PLACEHOLDER_FOR_SUBSTATION_NAME' in run.text:
                    run.text = run.text.replace('PLACEHOLDER_FOR_SUBSTATION_NAME', substation_name)
            break

    # Поиск таблицы в шаблоне
    table = template.tables[0]

    # Заполнение таблицы данными из запроса
    for record in queryset:
        row = table.add_row().cells

        # Преобразование времени в локальную временную зону
        local_time = record.pub_date.astimezone(local_timezone)

        if record.emergency_event and not record.short_circuit:
            cell = row[0]
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FF0000"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            paragraph = row[0].paragraphs[0]
            run = paragraph.add_run(local_time.strftime('%d.%m.%Y\n%H:%M'))
            font = run.font
            font.color.rgb = RGBColor(255, 255, 255)  # Белый цвет текста
            font.underline = True
            run.underline = WD_UNDERLINE.WAVY
        elif not record.emergency_event and record.short_circuit:
            cell = row[0]
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0000FF"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            paragraph = row[0].paragraphs[0]
            run = paragraph.add_run(local_time.strftime('%d.%m.%Y\n%H:%M'))
            font = run.font
            font.color.rgb = RGBColor(255, 255, 255)  # Белый цвет текста
            font.underline = True
            run.underline = WD_UNDERLINE.WAVY
        elif record.emergency_event and record.short_circuit:
            cell = row[0]
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="ffff00"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            paragraph = row[0].paragraphs[0]
            run = paragraph.add_run(local_time.strftime('%d.%m.%Y\n%H:%M'))
            font = run.font
            font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет текста
            font.underline = True
            run.underline = WD_UNDERLINE.WAVY
        else:
            row[0].text = local_time.strftime('%d.%m.%Y\n%H:%M')

        if record.entry_is_valid:
            if record.special_regime_introduced:
                paragraph = row[1].paragraphs[0]
                run = paragraph.add_run(record.text)
                font = run.font
                font.color.rgb = RGBColor(255, 0, 0)
                paragraph.add_run(f'\n{record.user.position} {record.user}')
            else:
                row[1].text = f'{record.text}\n{record.user.position} {record.user}'
        else:
            paragraph = row[1].paragraphs[0]
            run = paragraph.add_run('ОШИБОЧНАЯ ЗАПИСЬ!\n')
            font = run.font
            font.strike = False
            run = paragraph.add_run(record.text)
            font = run.font
            font.strike = True
            paragraph.add_run(f'\n{record.user.position} {record.user}')

        row[2].text = f'{record.comment.text}\n{record.comment.user.position} {record.comment.user}' if record.comment else ""

        # Вызов CoInitialize
    pythoncom.CoInitialize()

        # Сохранение документа во временный файл Word
    temp_docx_path = 'temp.docx'
    template.save(temp_docx_path)

        # Конвертация Word-документа в PDF
    temp_pdf_path = 'temp.pdf'
    docx2pdf.convert(temp_docx_path, temp_pdf_path)

        # Открытие PDF-файла и чтение его содержимого
    with open(temp_pdf_path, 'rb') as f:
        pdf_content = f.read()

        # Удаление временных файлов
    os.remove(temp_docx_path)
    os.remove(temp_pdf_path)

    # Установка заголовков HTTP-ответа для PDF-файла
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=exported_records.pdf'

        # Запись содержимого PDF в ответ
    response.write(pdf_content)

    return response